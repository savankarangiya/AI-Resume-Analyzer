import streamlit as st
from pypdf import PdfReader
import re
import os
from dotenv import load_dotenv
from groq import Groq
from docx import Document
from io import BytesIO
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

load_dotenv()

st.set_page_config(
    page_title="AI Career Copilot",
    layout="wide",
    initial_sidebar_state="collapsed"
)


# ------------ CUSTOM CSS ------------
st.markdown("""
<style>

.loader-overlay {
    position: fixed;
    inset: 0;
    background: rgba(255, 255, 255, 0.65);
    backdrop-filter: blur(8px);
    z-index: 999999;
    display: flex;
    align-items: center;
    justify-content: center;
}

.loader-card {
    background: white;
    padding: 34px 42px;
    border-radius: 24px;
    box-shadow: 0 25px 60px rgba(79, 70, 229, 0.30);
    text-align: center;
    border: 1px solid #e0e7ff;
    min-width: 340px;
}

.loader-spinner {
    width: 64px;
    height: 64px;
    border: 6px solid #e0e7ff;
    border-top: 6px solid #4f46e5;
    border-radius: 50%;
    animation: spin 0.85s linear infinite;
    margin: 0 auto 18px auto;
}

.loader-title {
    font-size: 22px;
    font-weight: 800;
    color: #312e81;
    margin-bottom: 8px;
}

.loader-subtitle {
    font-size: 15px;
    color: #64748b;
}

.block-container {
    padding-top: 2rem;
    padding-bottom: 3rem;
    max-width: 1400px;
    background-color: white;
    color: #0f172a;
}

@keyframes spin {
    0% { transform: rotate(0deg); }
    100% { transform: rotate(360deg); }
}

[data-testid="stSidebar"] {
    background: #f8fafc;
    border-right: 1px solid #e2e8f0;
}

div.stButton > button[kind="secondary"] {
    background: #eef2ff;
    color: #4f46e5;
    border: 1px solid #c7d2fe;
    border: none;
    border-radius: 14px;
    padding: 0.8rem 1.2rem;
    font-weight: 700;
    box-shadow: 0 10px 25px rgba(37, 99, 235, 0.25);
}

div.stButton > button {
    width: 100%;
    border-radius: 12px;
    border: 1px solid #dbeafe;
    background: white;
    color: #1e293b;
    font-weight: 600;
    padding: 0.7rem 1rem;
    transition: all 0.25s ease;
}

div.stButton > button:hover {
    background: linear-gradient(135deg, #667eea, #764ba2);
    color: white;
    border-color: transparent;
    transform: translateY(-2px);
    box-shadow: 0 8px 20px rgba(102, 126, 234, 0.35);
}

div[data-testid="stFileUploader"] {
    background: linear-gradient(135deg,#f8fafc,#eef2ff);
    border: 2px dashed #818cf8;
    border-radius: 22px;
    padding: 2rem;
    min-height: 150px;
    box-shadow: 0 12px 30px rgba(79,70,229,0.10);
}

div[data-testid="stFileUploader"]:hover {
    border-color: #4f46e5;
    background: linear-gradient(135deg,#eef2ff,#f5f3ff);
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div style="
padding:25px;
border-radius:20px;
background:linear-gradient(135deg,#667eea,#764ba2);
text-align:center;
margin-bottom:20px;
box-shadow:0 8px 25px rgba(0,0,0,0.2);
">
    <h1 style="color:white;margin:0;">
        🚀 AI Career Copilot
    </h1>
    <p style="color:white;font-size:18px;">
        Smart Resume Analysis Powered by AI
    </p>
</div>
""", unsafe_allow_html=True)


left_col, right_col = st.columns(2)

with left_col:
    st.subheader("📄 Upload Resume")

    uploaded_file = st.file_uploader(
        "Upload your resume in PDF format",
        type=["pdf"],
        help="Only text-based PDF resumes are supported."
    )

with right_col:
    st.subheader("💼 Job Description")

    job_description = st.text_area(
        "Paste Job Description Here",
        height=200,
        placeholder="Paste the internship or job description here..."
    )


jd_button_area = st.empty()
jd_match_clicked = False
jd_warning_area = st.empty()

if uploaded_file is not None:
    current_file_name = uploaded_file.name
    current_file_size = uploaded_file.size

    current_file_id = f"{current_file_name}_{current_file_size}"

    file_size_kb = round(current_file_size / 1024, 1)

    st.markdown(
    f"""
    <div style="background:white;padding:18px;border-radius:18px;border:1px solid #dbeafe;box-shadow:0 8px 24px rgba(37,99,235,0.08);margin-top:15px;margin-bottom:20px;">
        <h4 style="margin:0;color:#1e293b;">📄 {current_file_name}</h4>
        <p style="margin-top:8px;margin-bottom:0;color:#64748b;font-size:14px;">
            File Size: <b>{file_size_kb} KB</b> • Status: <b style="color:#16a34a;">Ready for AI Analysis</b>
        </p>
    </div>
    """,
    unsafe_allow_html=True
)

    if "last_file_id" not in st.session_state:
        st.session_state.last_file_id = current_file_id

    if st.session_state.last_file_id != current_file_id:
        st.session_state.project_result = None
        st.session_state.ats_result = None
        st.session_state.missing_result = None
        st.session_state.career_result = None
        st.session_state.optimized_resume = None
        st.session_state.resume_issues = None
        st.session_state.jd_match_result = None
        st.session_state.resume_text = ""
        st.session_state.resume_text_file = None
        st.session_state.last_jd = ""
        st.session_state.last_file_id = current_file_id

if uploaded_file is None:
    st.info(
        "🚀 Upload your resume to unlock ATS Score, Resume Optimization, Missing Skills Analysis, Project Extraction and JD Matching."
    )

# if uploaded_file:
#     st.session_state.project_result = None

if "project_result" not in st.session_state:
    st.session_state.project_result = None

if "ats_result" not in st.session_state:
    st.session_state.ats_result = None

if "missing_result" not in st.session_state:
    st.session_state.missing_result = None

if "career_result" not in st.session_state:
    st.session_state.career_result = None

if "optimized_resume" not in st.session_state:
    st.session_state.optimized_resume = None

if "resume_issues" not in st.session_state:
    st.session_state.resume_issues = None

if "jd_match_result" not in st.session_state:
    st.session_state.jd_match_result = None


GROQ_API_KEY = os.getenv("GROQ_API_KEY") 

def get_groq_client():
    if not GROQ_API_KEY:
        st.error("GROQ_API_KEY missing hai. .env file check karo.")
        st.stop()

    return Groq(api_key=GROQ_API_KEY)

def create_docx_from_text(resume_text):
    doc = Document()

    for line in resume_text.split("\n"):
        clean_line = line.strip()

        if not clean_line:
            continue

        clean_line = clean_line.replace("**", "")

        if clean_line.startswith("###"):
            heading_text = clean_line.replace("###", "").strip()
            doc.add_heading(heading_text, level=2)

        elif clean_line.startswith("#"):
            heading_text = clean_line.replace("#", "").strip()
            doc.add_heading(heading_text, level=1)

        elif clean_line.startswith("-"):
            bullet_text = clean_line.replace("-", "").strip()
            doc.add_paragraph(bullet_text, style="List Bullet")

        else:
            doc.add_paragraph(clean_line)

    docx_file = BytesIO()
    doc.save(docx_file)
    docx_file.seek(0)

    return docx_file


def extract_urls(text):
    raw_urls = re.findall(
        r'(?:https?://|www\.)[A-Za-z0-9\-._~:/?#\[\]@!$&\'()*+,;=%]+',
        text
    )

    cleaned_urls = []
    section_words = [
        "Professional", "Summary", "Education", "Technical", "Skills",
        "Projects", "Project", "Email", "Soft", "Contact", "Links"
    ]

    for url in raw_urls:
        clean_url = url.strip().strip(".,;:)]}")

        # PDF text extraction sometimes joins the next section/title with a URL.
        # Example: github.com/user/repoEmail -> github.com/user/repo
        for word in section_words:
            if clean_url.endswith(word) and len(clean_url) > len(word) + 10:
                clean_url = clean_url[:-len(word)]

        if clean_url and clean_url not in cleaned_urls:
            cleaned_urls.append(clean_url)

    return cleaned_urls


def extract_emails(text):
    return re.findall(r'\b[\w\.-]+@[\w\.-]+\.\w+\b', text)


def extract_phone_numbers(text):
    return re.findall(r'\+?\d[\d\s-]{8,}\d', text)


def extract_contact_block(text):
    lines = text.split("\n")
    contact_lines = []

    for line in lines[:10]:
        clean_line = line.strip()

        if not clean_line:
            continue

        if "professional summary" in clean_line.lower():
            break

        contact_lines.append(clean_line)

    return "\n".join(contact_lines).strip()


def normalize_token(text):
    return re.sub(r'[^a-z0-9]', '', text.lower())


def extract_protected_terms(text):
    protected_terms = []

    # Preserve hyphenated original words exactly: Scikit-learn, TF-IDF,
    # desktop-based, category-wise, non-spam, AI-Powered, etc.
    hyphenated_terms = re.findall(
        r'\b[A-Za-z0-9]+(?:-[A-Za-z0-9]+)+\b',
        text
    )

    # Preserve common technical names exactly when they are present in the original resume.
    known_terms = [
        "Machine Learning", "Deep Learning", "Artificial Intelligence",
        "Scikit-learn", "TF-IDF", "Naive Bayes", "Logistic Regression",
        "Python", "Java", "MySQL", "NumPy", "Pandas", "Matplotlib",
        "Tkinter", "GitHub", "Git", "VS Code", "Canva", "AWS",
        "Google Cloud", "TensorFlow", "PyTorch", "FastAPI", "Docker",
        "React", "JavaScript", "HTML", "CSS"
    ]

    lower_text = text.lower()
    for term in known_terms:
        if term.lower() in lower_text:
            protected_terms.append(term)

    protected_terms.extend(hyphenated_terms)

    unique_terms = []
    seen = set()
    for term in protected_terms:
        key = normalize_token(term)
        if key and key not in seen:
            unique_terms.append(term)
            seen.add(key)

    # Longer terms first avoids partial replacement problems.
    return sorted(unique_terms, key=len, reverse=True)


def preserve_original_terms(optimized_text, original_text):
    for term in extract_protected_terms(original_text):
        compact = normalize_token(term)

        if not compact:
            continue

        # Fix common AI formatting damage by replacing compact versions with original form.
        # Examples: Scikitlearn -> Scikit-learn, TFIDF -> TF-IDF,
        # desktopbased -> desktop-based, nonspam -> non-spam.
        optimized_text = re.sub(
            rf'\b{re.escape(compact)}\b',
            term,
            optimized_text,
            flags=re.IGNORECASE
        )

    return optimized_text


def remove_generated_links_sections(text):
    # Remove AI-created duplicate sections that often contain rewritten/wrong links.
    text = re.sub(
        r'\n\s*(Links|Contact Information)\s*\n[\s\S]*$',
        '',
        text,
        flags=re.IGNORECASE
    )
    return text.strip()


def remove_unverified_generated_urls(optimized_text, original_urls):
    original_url_keys = {normalize_token(url) for url in original_urls}

    cleaned_lines = []
    for line in optimized_text.split("\n"):
        line_urls = extract_urls(line)

        if line_urls:
            has_unverified_url = any(
                normalize_token(url) not in original_url_keys
                for url in line_urls
            )

            # Remove only lines that contain AI-invented or rewritten URLs.
            if has_unverified_url:
                continue

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()


def prepend_original_contact_block(optimized_text, original_contact_block):
    if not original_contact_block:
        return optimized_text.strip()

    # Remove the same contact lines wherever AI placed them, then force them to the top.
    cleaned_text = optimized_text

    for line in original_contact_block.split("\n"):
        line = line.strip()
        if line:
            cleaned_text = re.sub(
                rf'^\s*{re.escape(line)}\s*$',
                '',
                cleaned_text,
                flags=re.MULTILINE
            )

    cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text).strip()

    return original_contact_block.strip() + "\n\n" + cleaned_text


def finalize_optimized_resume(original_text, optimized_text):
    original_urls = extract_urls(original_text)
    original_contact_block = extract_contact_block(original_text)

    optimized_text = optimized_text.strip()
    optimized_text = remove_generated_links_sections(optimized_text)
    

    lines = optimized_text.split("\n")
    cleaned_lines = []
    seen = set()

    for line in lines:
        normalized = line.strip().lower()

        if normalized and normalized in seen:
            continue    

        seen.add(normalized)
        cleaned_lines.append(line)

    optimized_text = "\n".join(cleaned_lines)

    optimized_text = remove_unverified_generated_urls(optimized_text, original_urls)
    optimized_text = preserve_original_terms(optimized_text, original_text)
    optimized_text = prepend_original_contact_block(optimized_text, original_contact_block)
    # optimized_text = remove_duplicate_contact_lines(
    #     optimized_text,
    #     original_contact_block
    # )
    # optimized_text = append_missing_original_urls(optimized_text, original_urls)

    return optimized_text.strip()

def create_pdf_from_text(resume_text):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.pagesizes import A4

    pdf_file = BytesIO()

    doc = SimpleDocTemplate(
        pdf_file,
        pagesize=A4
    )

    styles = getSampleStyleSheet()

    content = []

    for line in resume_text.split("\n"):

        line = line.strip()

        line = line.replace("**", "")

        if not line:
            content.append(Spacer(1, 8))
            continue

        if line.startswith("#"):
            content.append(
                Paragraph(
                    line.replace("#", "").strip(),
                    styles["Heading2"]
                )
            )
        else:
            if line.startswith("###"):
                content.append(
                    Paragraph(line.replace("###", "").strip(), styles["Heading3"])
                )

            elif line.startswith("#"):
                content.append(
                    Paragraph(line.replace("#", "").strip(), styles["Heading2"])
                )

            elif line.startswith("-"):
                content.append(
                    Paragraph("• " + line.replace("-", "").strip(), styles["BodyText"])
                )

            else:
                content.append(
                    Paragraph(line, styles["BodyText"])
                )

    doc.build(content)

    pdf_file.seek(0)

    return pdf_file


def calculate_resume_score(resume_text):

    skills_db = [
        "Python", "Java", "MySQL", "React", "HTML", "CSS",
        "JavaScript", "C++", "C",
        "Machine Learning", "Deep Learning",
        "TensorFlow", "PyTorch",
        "Pandas", "NumPy", "Scikit-learn",
        "FastAPI", "Git", "GitHub", "Docker"
    ]

    found_skills = []

    for skill in skills_db:
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, resume_text.lower()):
            found_skills.append(skill)

    github_found = "github.com" in resume_text.lower()
    linkedin_found = "linkedin.com" in resume_text.lower()

    project_keywords = [
        "project",
        "developed",
        "built",
        "implemented",
        "integrated",
        "created",
        "designed"
    ]

    project_count = sum(
        resume_text.lower().count(keyword)
        for keyword in project_keywords
    )

    education_found = any(
        word in resume_text.lower()
        for word in [
            "b.tech", "b.e", "bca", "mca",
            "b.voc", "bsc", "msc", "degree"
        ]
    )

    email_found = "@" in resume_text

    phone_found = any(char.isdigit() for char in resume_text)

    certification_found = (
        "certification" in resume_text.lower()
        or "certificate" in resume_text.lower()
    )

    experience_found = (
        "experience" in resume_text.lower()
        or "internship" in resume_text.lower()
    )

    skills_score = min(len(found_skills) * 3, 30)

    project_score = min(project_count * 8, 30)

    education_score = 15 if education_found else 0

    completeness_score = 0

    if email_found:
        completeness_score += 3

    if phone_found:
        completeness_score += 3

    if github_found:
        completeness_score += 3

    if linkedin_found:
        completeness_score += 3

    if len(found_skills) >= 5:
        completeness_score += 3

    improvement_score = 0

    if certification_found:
        improvement_score += 5

    if experience_found:
        improvement_score += 5

    final_score = (
        skills_score
        + project_score
        + education_score
        + completeness_score
        + improvement_score
    )

    return min(final_score, 100)

def calculate_jd_match_score(resume_text, job_description):
    resume_words = set(re.findall(r'\b[a-zA-Z][a-zA-Z+#.-]*\b', resume_text.lower()))
    jd_words = set(re.findall(r'\b[a-zA-Z][a-zA-Z+#.-]*\b', job_description.lower()))

    stop_words = {
        "the", "and", "or", "to", "of", "in", "for", "with", "a", "an",
        "is", "are", "as", "on", "by", "from", "this", "that", "will",
        "be", "you", "your", "we", "our", "at", "it", "role", "job"
    }

    jd_keywords = jd_words - stop_words

    if len(jd_keywords) < 8:
        return 0, [], []

    matching_keywords = sorted(list(jd_keywords.intersection(resume_words)))
    missing_keywords = sorted(list(jd_keywords - resume_words))

    match_score = int((len(matching_keywords) / len(jd_keywords)) * 100)

    return match_score, matching_keywords[:15], missing_keywords[:15]


def expand_short_jd(job_description):
    jd = job_description.strip()

    if len(jd.split()) >= 25:
        return jd

    return f"""
    Job Role / Requirement:
    {jd}

    Expected Skills:
    - Python
    - Machine Learning
    - Data Analysis
    - Pandas
    - NumPy
    - Scikit-learn
    - Git
    - GitHub
    - Problem Solving
    - Basic project experience

    Responsibilities:
    - Build and understand Python-based applications
    - Work with machine learning models
    - Analyze data and extract insights
    - Use libraries like Pandas, NumPy, and Scikit-learn
    - Maintain code using Git and GitHub

    Candidate Expectations:
    - Fresher or internship-level candidate
    - Good technical foundation
    - Practical projects
    - Clean resume with relevant skills and links
    """


def show_ai_loader(title, subtitle):
    st.markdown(f"""
    <div class="loader-overlay">
        <div class="loader-card">
            <div class="loader-spinner"></div>
            <div class="loader-title">{title}</div>
            <div class="loader-subtitle">{subtitle}</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

def start_loader(title, subtitle):
    loader_placeholder = st.empty()

    with loader_placeholder:
        show_ai_loader(title, subtitle)

    return loader_placeholder

def safe_extract_section(text, start_heading, end_heading=None):
    try:
        if start_heading not in text:
            return "Not available"

        section = text.split(start_heading, 1)[1]

        if end_heading and end_heading in section:
            section = section.split(end_heading, 1)[0]

        return section.strip()

    except Exception:
        return "Not available"

if uploaded_file is not None:
    st.sidebar.markdown("""
    <div style="
    background:linear-gradient(135deg,#4f46e5,#7c3aed);
    padding:18px;
    border-radius:18px;
    text-align:center;
    margin-bottom:18px;
    box-shadow:0 8px 22px rgba(79,70,229,0.25);
    ">
    <h3 style="color:white;margin:0;">
    🤖 AI Features
    </h3>
    <p style="
    color:#e0e7ff;
    font-size:13px;
    margin-top:6px;
    margin-bottom:0;
    ">
    Choose an AI action
    </p>
    </div>
    """, unsafe_allow_html=True)


    st.markdown("""
    <div style="
    background:linear-gradient(135deg,#dcfce7,#ecfdf5);
    padding:18px;
    border-radius:16px;
    border:1px solid #86efac;
    margin-top:10px;
    margin-bottom:20px;
    ">
    <h4 style="
    margin:0;
    color:#166534;
    ">
    ✅ Resume Uploaded Successfully
    </h4>

    <p style="
    margin-top:8px;
    margin-bottom:0;
    color:#166534;
    ">
    Your resume has been processed and AI analysis features are now available.
    </p>
    </div>
    """, unsafe_allow_html=True)

    # ---------- PDF READ ----------
    if "resume_text" not in st.session_state:
        st.session_state.resume_text = ""

    if (
        st.session_state.resume_text == ""
        or st.session_state.get("resume_text_file") != current_file_id
    ):
        reader = PdfReader(uploaded_file)

        all_text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                all_text += page_text + "\n"

        st.session_state.resume_text = all_text
        st.session_state.resume_text_file = current_file_id
    else:
        all_text = st.session_state.resume_text

    if not all_text.strip():
        st.error("Unable to extract text from this PDF. Please upload a text-based PDF.")
        st.stop()


    # ---------- SKILLS ----------
    skills_db = [
        "Python", "Java", "MySQL", "React", "HTML", "CSS",
        "JavaScript", "C++", "C",
        "Machine Learning", "Deep Learning",
        "TensorFlow", "PyTorch",
        "Pandas", "NumPy", "Scikit-learn",
        "FastAPI", "Git", "GitHub", "Docker"
    ]

    found_skills = []

    for skill in skills_db:
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, all_text.lower()):
            found_skills.append(skill)

    st.markdown(f"""
    <div style="
    background:white;
    padding:18px;
    border-radius:18px;
    border:1px solid #e5e7eb;
    box-shadow:0 6px 18px rgba(0,0,0,0.05);
    margin-bottom:15px;
    ">
    <h2 style="margin:0;">🧠 Detected Skills</h2>

    <p style="
    margin-top:8px;
    color:#64748b;
    font-size:15px;
    ">
    Total Skills Detected: <b>{len(found_skills)}</b>
    </p>
    </div>
    """, unsafe_allow_html=True)

    cols = st.columns(4)

    for index, skill in enumerate(found_skills):
        with cols[index % 4]:
            st.markdown("<br>", unsafe_allow_html=True)
            st.success(f"✅ {skill}")

    # ---------- RESUME METRICS ----------

    github_found = "github.com" in all_text.lower()

    linkedin_found = "linkedin.com" in all_text.lower()

    project_count = all_text.lower().count("project")

    education_found = any(
        word in all_text.lower()
        for word in [
            "b.tech",
            "b.e",
            "bca",
            "mca",
            "b.voc",
            "bsc",
            "msc",
            "degree"
        ]
    )


    # --------------- Contact info detection ----------------

    email_found = "@" in all_text

    phone_found = any(char.isdigit() for char in all_text)

    certification_found = (
        "certification" in all_text.lower()
        or "certificate" in all_text.lower()
    )

    experience_found = (
        "experience" in all_text.lower()
        or "internship" in all_text.lower()
    )

    # ---------- RULE BASED ATS SCORE ----------

    skills_score = min(len(found_skills) * 3, 30)

    project_score = min(project_count * 8, 30)

    education_score = 15 if education_found else 0

    completeness_score = 0

    if email_found:
        completeness_score += 3

    if phone_found:
        completeness_score += 3

    if github_found:
        completeness_score += 3

    if linkedin_found:
        completeness_score += 3

    if len(found_skills) >= 5:
        completeness_score += 3

    improvement_score = 0

    if certification_found:
        improvement_score += 5

    if experience_found:
        improvement_score += 5

    rule_based_ats_score = calculate_resume_score(all_text)


    # ---------- GROQ ATS SCORE ----------

    st.sidebar.markdown("### 📊 ATS & Matching")

    if st.sidebar.button("📊 Generate ATS Score"):
        client = get_groq_client()

        loader_placeholder = st.empty()

        with loader_placeholder:
            show_ai_loader(
                "🤖 AI ATS Analysis Running",
                "Evaluating skills, projects, resume structure and ATS compatibility..."
            )

        with st.spinner(""):

            ats_prompt = f"""
            You are an ATS resume evaluator for an entry-level AI/ML internship role.

            The calculated ATS score is {rule_based_ats_score}/100.

            You must use this exact score.
            Do not create a different score.

            Evaluate this resume strictly out of 100 using this scoring system:

            1. Technical Skills Match: 30 marks
            - Python, Machine Learning, Pandas, NumPy, Scikit-learn, SQL, Git/GitHub

            2. Project Quality: 30 marks
            - AI/ML projects
            - clear project descriptions
            - technologies mentioned
            - GitHub/project links
            - measurable or practical impact

            3. Education & Academic Fit: 15 marks
            - IT/CS/AI/ML related education
            - CGPA or percentage mentioned

            4. Resume Completeness: 15 marks
            - contact info
            - GitHub/LinkedIn
            - skills section
            - project section
            - clean structure

            5. Improvement Potential: 10 marks
            - missing deployment
            - missing internship
            - missing certifications
            - missing advanced AI/ML skills

            Return exactly in this format:

            ## ATS Score
            <score>/100

            ## Score Breakdown
            - Technical Skills Match: <marks>/30
            - Project Quality: <marks>/30
            - Education & Academic Fit: <marks>/15
            - Resume Completeness: <marks>/15
            - Improvement Potential: <marks>/10

            ## Strengths
            - point 1
            - point 2
            - point 3

            ## Weaknesses
            - point 1
            - point 2

            ## Top Improvements
            - point 1
            - point 2
            - point 3

            Be strict. Do not give inflated scores.

            Important Rules:

            - If the resume is empty, give ATS Score: 0/100
            - If no projects are present, Project Quality must be below 10/30
            - If no technical skills are present, Technical Skills Match must be below 5/30
            - If GitHub and LinkedIn are missing, Resume Completeness cannot exceed 8/15
            - Do not assume information that is not present in the resume
            - Penalize missing internships, certifications, projects, and measurable achievements
            - Freshers without strong projects should rarely score above 70/100

            Resume:
            {all_text}
            """
            try:

                if st.session_state.ats_result is None:

                    ats_response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "user", "content": ats_prompt}
                        ],
                        temperature=0,
                    )

                    st.session_state.ats_result = (
                        ats_response.choices[0].message.content
                    )

                ats_text = st.session_state.ats_result

                loader_placeholder.empty()

                if rule_based_ats_score >= 80:
                    score_status = "Excellent"
                    score_color = "#16a34a"

                elif rule_based_ats_score >= 70:
                    score_status = "Good"
                    score_color = "#2563eb"

                elif rule_based_ats_score >= 60:
                    score_status = "Average"
                    score_color = "#f59e0b"

                else:
                    score_status = "Needs Improvement"
                    score_color = "#dc2626"

                left_col, right_col = st.columns([1, 2])

                with left_col:
                    st.markdown(f"""
                    <div style="
                        background:white;
                        padding:24px;
                        border-radius:20px;
                        box-shadow:0 8px 24px rgba(0,0,0,0.08);
                        border:1px solid #e2e8f0;
                        text-align:center;
                    ">
                        <h3 style="margin-bottom:10px;color:#0f172a;">Your Score</h3>
                        <h1 style="font-size:46px;margin:0;color:{score_color};">{rule_based_ats_score}/100</h1>
                        <p style="color:#64748b;margin-top:8px;">ATS Resume Score</p>
                        <p style="
                        font-size:14px;
                        font-weight:600;
                        color:{score_color};
                        margin-top:5px;
                        ">
                        {score_status}
                        </p>
                        <div style="
                        width:100%;
                        height:12px;
                        background:#e5e7eb;
                        border-radius:999px;
                        margin-top:15px;
                        overflow:hidden;
                        ">
                            <div style="
                            width:{rule_based_ats_score}%;
                            height:100%;
                            background:{score_color};
                            border-radius:999px;
                            ">
                            </div>
                        </div>
                        </div>
                    """, unsafe_allow_html=True)

                with right_col:
                    st.markdown("""
                    ...
                    """, unsafe_allow_html=True)


                    score_breakdown = safe_extract_section(
                        ats_text,
                        "## Score Breakdown",
                        "## Strengths"
                    )

                    strengths = safe_extract_section(
                        ats_text,
                        "## Strengths",
                        "## Weaknesses"
                    )

                    weaknesses = safe_extract_section(
                        ats_text,
                        "## Weaknesses",
                        "## Top Improvements"
                    )

                    improvements = safe_extract_section(
                        ats_text,
                        "## Top Improvements"
                    )

                    with st.container(border=True):

                        st.markdown("""
                        <style>
                        ...
                        </style>
                        """, unsafe_allow_html=True)

                        st.markdown("#### 📌 Score Breakdown")
                        st.markdown(score_breakdown)

                        st.markdown("#### ✅ Strengths")
                        st.success(strengths)

                        st.markdown("#### ⚠️ Weaknesses")
                        st.warning(weaknesses)

                        st.markdown("#### 🚀 Top Improvements")
                        st.info(improvements)

            except Exception as e:
                st.error(f"Error: {e}")

    # ---------- AI MISSING SKILLS ----------

    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🧠 Resume Analysis")

    if st.sidebar.button("🧩 Analyze Missing Skills"):
        client = get_groq_client()

        loader_placeholder = st.empty()
        with loader_placeholder:
            show_ai_loader(
                "🧩 AI Missing Skills Analysis",
                "Finding technical skills, tools and concepts missing from your resume..."
            )

        with st.spinner(""):
            missing_skill_prompt = f"""
            Analyze this resume for an AI/ML internship role.

            Find missing skills.

            Return exactly this format:

            Missing Technical Skills:
            - skill 1
            - skill 2

            Missing Tools:
            - tool 1
            - tool 2

            Missing Concepts:
            - concept 1
            - concept 2

            Resume:
            {all_text}
            """

            if st.session_state.missing_result is None:

                missing_response = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[
                        {"role": "user", "content": missing_skill_prompt}
                    ]
                )

                st.session_state.missing_result = (
                    missing_response.choices[0].message.content
                )

            missing_text = st.session_state.missing_result
            loader_placeholder.empty()

        st.subheader("🧩 AI Missing Skills Analysis")

        with st.container(border=True):
            st.markdown(missing_text)

    # ---------- CAREER RECOMMENDATIONS ----------

    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🚀 Resume Improvement")

    if st.sidebar.button("🎯 Career Recommendations"):
        client = get_groq_client()

        loader_placeholder = st.empty()
        with loader_placeholder:
            show_ai_loader(
                "🎯 AI Career Recommendations",
                "Finding suitable career roles based on your resume profile..."
            )

        with st.spinner(""):
            career_prompt = f"""
            Analyze this resume and suggest the top 5 suitable career roles.

            Return ONLY in this clean format:

            ## 1. Role Name: <role>

            **Why Suitable:**
            - <short reason>
            - <short reason>

            **Resume Evidence:**
            - <specific evidence from resume>

            ---

            ## 2. Role Name: <role>

            **Why Suitable:**
            - <short reason>
            - <short reason>

            **Resume Evidence:**
            - <specific evidence from resume>

            ---

            Rules:
            - Do not write introduction.
            - Do not write long paragraphs.
            - Use short bullet points.
            - Keep each role under 5 lines.

            Resume:
            {all_text}
            """

            try:

                if st.session_state.career_result is None:

                    career_response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "user", "content": career_prompt}
                        ]
                    )

                    st.session_state.career_result = (
                        career_response.choices[0].message.content
                    )

                career_text = st.session_state.career_result
                loader_placeholder.empty()

                st.subheader("🎯 Career Recommendations")

                with st.container(border=True):
                    st.markdown(career_text)

            except Exception as e:
                st.error(f"Error: {e}")

    # ---------- AI RESUME ISSUES DETECTOR ----------

    if st.sidebar.button("🔍 Find Resume Issues"):
        client = get_groq_client()

        loader_placeholder = st.empty()
        with loader_placeholder:
            show_ai_loader(
                "🔍 AI Resume Issues Detector",
                "Checking weak sections, missing proof, ATS gaps and improvement areas..."
            )

        with st.spinner(""):
            issues_prompt = f"""
            You are a strict resume reviewer for AI/ML internship and fresher roles.

            Analyze this resume and find real problems.

            Return ONLY in this format:

            ## 🔍 Resume Issues Found

            - Issue 1
            - Issue 2
            - Issue 3
            - Issue 4
            - Issue 5

            ## 💡 Recommended Fixes

            - Fix 1
            - Fix 2
            - Fix 3
            - Fix 4
            - Fix 5

            Rules:
            - Be strict.
            - Do not write long paragraphs.
            - Do not praise the resume.
            - Do not suggest fake experience.
            - Focus on ATS, projects, skills, measurable impact, internships, certifications, and structure.

            Resume:
            {all_text}
            """

            try:
                if st.session_state.resume_issues is None:

                    issues_response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "user", "content": issues_prompt}
                        ],
                        temperature=0,
                    )

                    st.session_state.resume_issues = (
                        issues_response.choices[0].message.content
                    )

                loader_placeholder.empty()

                st.markdown("## 🔍 AI Resume Issues Detector")

                with st.container(border=True):
                    st.markdown(st.session_state.resume_issues)

            except Exception as e:
                st.error(f"Error: {e}")

    # ---------- RESUME VS JOB DESCRIPTION MATCH ----------

    with jd_button_area.container():
        col_empty, col_button = st.columns([5, 1])

        with col_button:
            jd_match_clicked = st.button("🔍 Analyze JD", key="jd_match_button")

    if jd_match_clicked:
        client = get_groq_client()

        if not job_description.strip():
            jd_warning_area.warning("⚠️ Please paste a job description first.")
        else:
            loader_placeholder = st.empty()
            with loader_placeholder:
                show_ai_loader(
                    "💼 AI JD Match Analysis",
                    "Comparing your resume skills, projects and keywords with the job description..."
                )

            with st.spinner(""):

                expanded_jd = expand_short_jd(job_description)

                jd_score, matching_keywords, missing_keywords = calculate_jd_match_score(
                    all_text,
                    expanded_jd
               )

                if len(job_description.strip().split()) < 5:
                    loader_placeholder.empty()
                    st.error(
                        "⚠️ Please paste a detailed job description. Short text like 'Python development' is not enough for accurate matching."
                    )
                    jd_prompt = None
                else:

                    jd_prompt = f"""
                    You are an ATS and recruitment expert.

                    The calculated JD match score is {jd_score}/100.
                    You must use this exact score.

                    Compare the resume with the job description.

                    Return ONLY in this format:

                    ## Match Score
                    {jd_score}/100

                    ## Matching Skills
                    - skill 1
                    - skill 2
                    - skill 3

                    ## Missing Keywords
                    - keyword 1
                    - keyword 2
                    - keyword 3

                    ## Resume Gaps
                    - gap 1
                    - gap 2
                    - gap 3

                    ## How to Improve Match
                    - improvement 1
                    - improvement 2
                    - improvement 3

                    Rules:
                    - Be strict.
                    - Do not invent skills.
                    - Only compare resume content with JD.
                    - If JD asks for a skill missing from resume, put it in Missing Keywords.
                    - Keep bullets short and useful.

                    Resume:
                    {all_text}

                    Job Description:
                    {expanded_jd}
                    """

                if jd_prompt is not None:
                    try:
                        current_jd = job_description.strip()

                        if "last_jd" not in st.session_state:
                            st.session_state.last_jd = current_jd

                        if st.session_state.last_jd != current_jd:
                            st.session_state.jd_match_result = None
                            st.session_state.last_jd = current_jd

                        if st.session_state.jd_match_result is None:

                            jd_response = client.chat.completions.create(
                                model="llama-3.1-8b-instant",
                                messages=[
                                    {"role": "user", "content": jd_prompt}
                                ],
                                temperature=0,
                            )

                            st.session_state.jd_match_result = (
                                jd_response.choices[0].message.content
                            )

                        loader_placeholder.empty()

                        st.markdown("## 💼 Resume vs Job Description Match")

                        with st.container(border=True):
                            st.markdown(st.session_state.jd_match_result)

                    except Exception as e:
                        loader_placeholder.empty()
                        st.error(f"Error: {e}")


    # ---------- AI RESUME OPTIMIZER ----------


    if st.sidebar.button("🛠 Fix & Download Resume"):
        client = get_groq_client()

        loader_placeholder = st.empty()
        with loader_placeholder:
            show_ai_loader(
                "🛠 AI Resume Optimizer Running",
                "Improving wording, structure, ATS keywords and resume presentation..."
            )


        with st.spinner(""):

            optimize_prompt = f"""
            You are a senior resume writer and ATS optimization expert.

            Your task is to completely rewrite and improve this resume for AI/ML Internship and Fresher Software Roles.

            STRICT RULES:

            - Do NOT invent fake internships.
            - Do NOT invent fake certifications.
            - Do NOT invent fake projects.
            - Do NOT invent fake skills.
            - Never remove the candidate's name.
            - Preserve all technical keywords exactly as written.
            - Preserve all URLs exactly as written.
            - Do not rewrite URLs.
            - Do not shorten URLs.
            - Do not modify LinkedIn URLs.
            - Do not modify GitHub URLs.
            - Copy URLs exactly from the original resume.
            - Do not modify technology names.
            - Keep "Scikit-learn" exactly as "Scikit-learn".
            - Keep "TF-IDF" exactly as "TF-IDF".
            - Keep "Machine Learning" exactly as written.
            - Keep "GitHub" exactly as written.
            - Never remove phone number.
            - Never remove email address.
            - Never remove LinkedIn URL.
            - Never remove GitHub URL.
            - Keep contact information at the top of the resume.
            - Use ONLY information already present in the resume.

            CRITICAL PRESERVATION RULES:

            Copy the following items exactly character-by-character.

            Do not rewrite them.
            Do not reformat them.
            Do not simplify them.

            Preserve:
            - Full Name
            - Phone Number
            - Email Address
            - LinkedIn URL
            - GitHub URL
            - Project Repository URLs
            - Scikit-learn
            - TF-IDF

            If any URL or keyword is modified, the output is considered incorrect.

            IMPORTANT PRESERVATION RULES:

            - Preserve the candidate name exactly as provided.
            - Preserve phone number exactly as provided.
            - Preserve email address exactly as provided.
            - Preserve LinkedIn URL exactly as provided.
            - Preserve GitHub URL exactly as provided.
            - Preserve all project repository URLs exactly as provided.
            - Do not rewrite usernames, URLs, repository names, or profile links.
            - If contact information exists in the original resume, it must appear unchanged in the improved resume.

            IMPROVEMENTS REQUIRED:

            - Fix grammar and formatting.
            - Rewrite weak bullet points into strong, professional, achievement-focused bullet points.
            - Improve project descriptions with clear problem, solution, technologies used, and outcome.
            - Add practical impact where it is clearly supported by the resume.
            - Use strong action verbs such as Developed, Built, Implemented, Integrated, Designed, Analyzed, Improved, Managed.
            - Improve ATS keyword matching for AI/ML, Python, data analysis, Git, GitHub, and software development roles.
            - Keep the resume suitable for a fresher or internship-level candidate.
            - Do not exaggerate experience.
            - Do not add fake numbers, fake metrics, fake certifications, or fake internships.
            - Keep every bullet concise, clear, and recruiter-friendly.
            - Preserve all genuine information.

            PROJECT IMPROVEMENT RULES:

            For every project:

            - Mention technologies used.
            - Explain what was built.
            - Explain key functionality.
            - Explain the problem solved by the project.
            - Explain how the solution works.
            - Convert short descriptions into 3–5 professional resume bullet points.
            - Use strong action verbs.
            - Highlight technical implementation details.
            - Mention machine learning, AI, database, API, automation, analytics, GUI, or deployment details when they already exist in the original project.
            - Do not invent features that are not present in the original resume.

            PROFESSIONAL SUMMARY RULES:

            - Professional Summary must be 3 to 4 lines only.
            - Focus on the candidate's actual education, skills, and projects.
            - Do not use generic phrases such as:
            "Highly motivated"
            "Hardworking individual"
            "Seeking a challenging opportunity"
            "Passionate professional"
            - Make the summary specific to the resume content.
            - Mention AI/ML, Python, software development, or data analysis only if they exist in the original resume.
            - Do not invent experience.

            PROJECT BULLET RULES:

            - Every project must contain at least 3 bullet points.
            - Each bullet should start with a strong action verb.
            - Focus on implementation, technologies, functionality, and technical contribution.
            - Prefer technical details over generic descriptions.
            - Mention libraries, frameworks, databases, APIs, machine learning models, or tools if they exist in the original resume.
            - Do not add fake percentages, fake accuracy values, fake performance improvements, or fake business impact.
            - Do not write vague statements like:
            "Worked on a project"
            "Learned many technologies"
            "Improved skills"
            - Every bullet must communicate a concrete technical contribution.

            OUTPUT FORMAT:

            # Professional Summary

            # Technical Skills

            # Projects

            # Education

            # Certifications (if present)

            IMPORTANT:

            - Do NOT create a Links section.
            - Do NOT create a Contact Information section.
            - Contact details already exist at the top.
            - Do not repeat name, email, phone, LinkedIn, GitHub, or URLs anywhere else in the resume.
            - Every project description must use bullet points beginning with "-".

            Return ONLY the improved resume.
            Do NOT explain changes.
            Do NOT give suggestions.
            Do NOT compare with the original.

            Resume:

            {all_text}
            """
            try:
                if st.session_state.optimized_resume is None:


                    optimize_response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "user", "content": optimize_prompt}
                        ],
                        temperature=0.2,    
                    )

                    optimized_text = optimize_response.choices[0].message.content

                    optimized_text = finalize_optimized_resume(
                        all_text,
                        optimized_text
                    )

                    st.session_state.optimized_resume = optimized_text

                loader_placeholder.empty()

                original_score = calculate_resume_score(all_text)

                improved_score = calculate_resume_score(st.session_state.optimized_resume)

                score_difference = improved_score - original_score

                st.markdown("## 📈 Resume Improvement Score")

                col1, col2, col3 = st.columns(3)

                with col1:
                    st.metric("Original ATS Score", f"{original_score}/100")

                with col2:
                    st.metric("Improved ATS Score", f"{improved_score}/100")

                with col3:
                    st.metric("Improvement", f"{score_difference:+}/100")

                st.markdown("## 🛠 AI Improved Resume")
                
                with st.container(border=True):
                    st.markdown(st.session_state.optimized_resume)

                docx_file = create_docx_from_text(
                    st.session_state.optimized_resume
                )

                pdf_file = create_pdf_from_text(
                    st.session_state.optimized_resume
                )

                st.download_button(
                    label="⬇️ Download Resume (.DOCX)",
                    data=docx_file,
                    file_name="improved_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.download_button(
                    label="⬇️ Download Resume(.PDF)",
                    data=pdf_file,
                    file_name="improved_resume.pdf",
                    mime="application/pdf"
                )

            except Exception as e:
                st.error(f"Error: {e}")

    # ---------- PROJECTS (FIXED) ----------
    projects = []
    current_project = ""

    project_started = False

    for line in all_text.split("\n"):
        clean = line.strip()

        # start project section
        if "project" in clean.lower():
            project_started = True
            continue

        # stop at other sections
        if clean.lower() in ["education", "experience", "skills"]:
            project_started = False

        if project_started:
            if clean.startswith("•"):
                current_project += " " + clean
            else:
                if current_project:
                    projects.append(current_project.strip())
                    current_project = ""
                if "|" in clean:
                    current_project = clean

    # last flush
    if st.sidebar.button("📂 Extract Projects"):
        client = get_groq_client()

        loader_placeholder = st.empty()
        with loader_placeholder:
            show_ai_loader(
                "📂 AI Project Extraction",
                "Extracting project names, technologies, descriptions and key highlights..."
            )

        with st.spinner(""):
            prompt = f"""
            Analyze this resume and extract all projects.

            Return ONLY markdown format.

            For each project use:

            ### 🚀 Project Name

            **🛠 Technologies**
            - item

            **📄 Description**
            Short description in 2-3 lines.

            **⭐ Key Highlights**
            - highlight 1
            - highlight 2

            ---

            Do not write:
            "Based on the provided resume"

            Do not write introductions.

            Resume:
            {all_text}
            """

            try:

                if st.session_state.project_result is None:

                    project_response = client.chat.completions.create(
                        model="llama-3.1-8b-instant",
                        messages=[
                            {"role": "user", "content": prompt}
                        ]
                    )

                    st.session_state.project_result = (
                        project_response.choices[0].message.content
                    )


                loader_placeholder.empty()

                st.markdown("""
                <div style="
                background:white;
                padding:24px;
                border-radius:20px;
                box-shadow:0 8px 24px rgba(0,0,0,0.08);
                border:1px solid #e2e8f0;
                margin-top:10px;
                margin-bottom:20px;
                ">
                <h2 style="margin-top:0;">📂 AI Extracted Projects</h2>
                </div>
                """, unsafe_allow_html=True)

                st.markdown(st.session_state.project_result)

            except Exception as e:
                st.error(f"Error: {e}")